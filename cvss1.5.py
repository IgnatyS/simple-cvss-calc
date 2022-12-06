# Версия 1.5. Исправления
# Исправляем диаграммы [ ]
# Тултип на все кнопки параметров [ ]

import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from pandas import DataFrame
import tkinter as tk
import openpyxl
import docx


class Main(tk.Frame):

    baseentry = None

    temporalentry = None

    environmentalentry = None

    vul_list = []

    mainlistbox = []
    OpisanieVul = []
    Vendor = []
    Nazvanie = []
    Klass = []
    UrovenOpasnosti = []
    Meri = []

    sheet = None

    def __init__(self, root):
        super().__init__(root)

        # Первая линия описние и текст
        namelabel = tk.Label(text="Введите исследуемый компонент :")
        namelabel.place(x=30, y=20, width=200, height=25)
        self.nameentry = tk.Entry()
        self.nameentry.place(x=30, y=55, width=500, height=25)
        self.searchbutton = tk.Button(text="Найти", bg="white", command=self.insert)
        self.searchbutton.place(x=540, y=55, width=130, height=25)
        self.clearbutton = tk.Button(text="Очистить", bg="white", command=self.clear)
        self.clearbutton.place(x=540, y=90, width=130, height=25)
        self.savebutton = tk.Button(text="Сохранить", bg="white", command=self.SaveForm)
        self.savebutton.place(x=540, y=125, width=130, height=25)
        self.insertdbtn = tk.Button(text="Добавить диаграммы", bg="white", command=self.insertdiagram)
        self.insertdbtn.place(x=400, y=365, width=130, height=25)

        # линия листбокса
        self.listbox = tk.Listbox()
        self.listbox.bind('<<ListboxSelect>>', self.select_item)

        self.listbox.place(x=30, y=90, width=500, height=150)

        # линия базовых показателей
        baselabel = tk.Label(text="Базовый балл")
        baselabel.place(x=50, y=285, width=200, height=25)
        Main.baseentry = tk.Entry()
        Main.baseentry.place(x=250, y=285, width=25, height=25)
        self.base_button = tk.Button(text="Добавить", bg="white", command=self.open_base)
        self.base_button.place(x=295, y=285, width=85, height=25)

        # линия временных показателей
        temporallabel = tk.Label(text="Временный балл")
        temporallabel.place(x=50, y=325, width=200, height=25)
        Main.temporalentry = tk.Entry()
        Main.temporalentry.place(x=250, y=325, width=25, height=25)
        self.temporalbutton = tk.Button(text="Добавить", bg="white", command=self.open_temporal)
        self.temporalbutton.place(x=295, y=325, width=85, height=25)

        # линия контекстных показателей
        environmentallabel = tk.Label(text="Контекстный балл")
        environmentallabel.place(x=50, y=365, width=200, height=25)
        Main.environmentalentry = tk.Entry()
        Main.environmentalentry.place(x=250, y=365, width=25, height=25)
        self.environmentalbutton = tk.Button(text="Добавить", bg="white", command=self.open_environmental)
        self.environmentalbutton.place(x=295, y=365, width=85, height=25)

        # Правое поле парсинга ексель файла 1
        label2 = tk.Label(text="Описание уязвимости :")
        label2.place(x=730, y=20, width=150, height=25)
        self.textbox2 = tk.Text()
        self.textbox2.place(x=730, y=55, width=500, height=185)

        # Правое поле парсинга ексель файла 2
        label3 = tk.Label(text="Вендор ПО :")
        label3.place(x=730, y=250, width=85, height=25)
        self.textbox3 = tk.Text()
        self.textbox3.place(x=730, y=285, width=500, height=25)

        # Правое поле парсинга ексель файла 3
        label4 = tk.Label(text="Название ПО :")
        label4.place(x=730, y=320, width=85, height=25)
        self.textbox4 = tk.Text()
        self.textbox4.place(x=730, y=355, width=500, height=25)

        # Правое поле парсинга ексель файла 4
        label5 = tk.Label(text="Класс уязвимости :")
        label5.place(x=730, y=390, width=120, height=25)
        self.textbox5 = tk.Text()
        self.textbox5.place(x=730, y=425, width=500, height=25)

        # Правое поле парсинга ексель файла 5
        label6 = tk.Label(text="Уровень опасности по CVSSv.2 :")
        label6.place(x=730, y=460, width=190, height=25)
        self.textbox6 = tk.Text()
        self.textbox6.place(x=730, y=495, width=500, height=25)

        # Правое поле парсинга ексель файла 5
        label6 = tk.Label(text="Возможные меры по устранению уязвимости :")
        label6.place(x=730, y=530, width=270, height=25)
        self.textbox7 = tk.Text()
        self.textbox7.place(x=730, y=565, width=500, height=185)

        # search
        wb = openpyxl.load_workbook("vullist.xlsx")
        Main.sheet = wb['Sheet']

        for apple_row in Main.sheet['B4':'B1000']:
            for apple_Obj in apple_row:
                Main.vul_list.append(apple_Obj.value)

    def select_item(self, event):

        value = (self.listbox.get(self.listbox.curselection()))
        Main.mainlistbox = value
        index = int(Main.vul_list.index(value))
        main_index = index + 4

        self.textbox2.delete(1.0, tk.END)
        self.textbox2.insert(1.0, Main.sheet['C' + str(main_index)].value)
        Main.OpisanieVul = str(Main.sheet['C' + str(main_index)].value)

        self.textbox3.delete(1.0, tk.END)
        self.textbox3.insert(1.0, Main.sheet['D' + str(main_index)].value)
        Main.Vendor = str(Main.sheet['D' + str(main_index)].value)

        self.textbox4.delete(1.0, tk.END)
        self.textbox4.insert(1.0, Main.sheet['E' + str(main_index)].value)
        Main.Nazvanie = str(Main.sheet['E' + str(main_index)].value)

        self.textbox5.delete(1.0, tk.END)
        self.textbox5.insert(1.0, Main.sheet['I' + str(main_index)].value)
        Main.Klass = str(Main.sheet['I' + str(main_index)].value)

        self.textbox6.delete(1.0, tk.END)
        self.textbox6.insert(1.0, Main.sheet['M' + str(main_index)].value)
        Main.UrovenOpasnosti = str(Main.sheet['M' + str(main_index)].value)

        self.textbox7.delete(1.0, tk.END)
        self.textbox7.insert(1.0, Main.sheet['N' + str(main_index)].value)
        Main.Meri = str(Main.sheet['N' + str(main_index)].value)



    def insert(self):
        search = str(self.nameentry.get())
        sort_list = list(filter(lambda string: search in string, Main.vul_list))
        self.listbox.delete(0, tk.END)
        for item in sort_list:
            self.listbox.insert(tk.END, item)

    def insertdiagram(self):
        # Данные 1
        data1 = {'BaseMetrics': ['B', 'I', 'E'],
                 'BasePoint': [Base.BaseScore, Base.Impact, Base.Exploitability]
                 }
        df1 = DataFrame(data1, columns=['BaseMetrics', 'BasePoint'])
        # Данные 2
        data2 = {'TemporalMetrics': ['T'],
                 'TemporalPoint': [Temporal.TemporalScore]
                 }
        df2 = DataFrame(data2, columns=['TemporalMetrics', 'TemporalPoint'])
        # Данные 3
        data3 = {'EnviromentalMetrics': ['E', 'MI'],
                 'EnviromentalPoint': [Environmental.EnvironmentalScore, Environmental.ModifiedImpact]
                 }
        df3 = DataFrame(data3, columns=['EnviromentalMetrics', 'EnviromentalPoint'])

        # Графоний 1
        figure1 = plt.Figure()
        ax1 = figure1.add_subplot(111)
        figure1.set_size_inches(6, 5, forward=True)
        canvas1 = FigureCanvasTkAgg(figure1, root)
        canvas1.get_tk_widget().place(x=30, y=500, width=200, height=200)
        df1 = df1[['BaseMetrics', 'BasePoint']].groupby('BaseMetrics').sum()
        df1.plot(kind='bar', legend=True, ax=ax1)
        ax1.set_title('Базовая оценка')

        # Графоний 2
        figure2 = plt.Figure()
        ax2 = figure2.add_subplot(111)
        figure2.set_size_inches(6, 5, forward=True)
        canvas2 = FigureCanvasTkAgg(figure2, root)
        canvas2.get_tk_widget().place(x=250, y=500, width=200, height=200)
        df2 = df2[['TemporalMetrics', 'TemporalPoint']].groupby('TemporalMetrics').sum()
        df2.plot(kind='bar', legend=True, ax=ax2)
        ax2.set_title('Временная оценка')

        # Графоний 3
        figure3 = plt.Figure()
        ax3 = figure3.add_subplot(111)
        figure3.set_size_inches(6, 5, forward=True)
        canvas3 = FigureCanvasTkAgg(figure3, root)
        canvas3.get_tk_widget().place(x=470, y=500, width=200, height=200)
        df3 = df3[['EnviromentalMetrics', 'EnviromentalPoint']].groupby('EnviromentalMetrics').sum()
        df3.plot(kind='bar', legend=True, ax=ax3)
        ax3.set_title('Контекстная оценка')

    def clear(self):
        self.textbox2.delete(1.0, tk.END)
        self.textbox3.delete(1.0, tk.END)
        self.textbox4.delete(1.0, tk.END)
        self.textbox5.delete(1.0, tk.END)
        self.textbox6.delete(1.0, tk.END)
        self.textbox7.delete(1.0, tk.END)
        self.listbox.delete(0, tk.END)


# Методы класса мейн для вызова дочерних классов - окон
    def open_base(self):
        Base()

    def open_temporal(self):
        Temporal()

    def open_environmental(self):
        Environmental()

    def SaveForm(self):
        SaveForm()


class Base(tk.Toplevel):

    # Дефолтные значения класса БАЗА
    Scope = 0
    AttackComplexity = 0
    AttackVector = 0
    UserInteraction = 0
    PrivilegesRequired = 0
    Confidentiality = 0
    Integrity = 0
    Availability = 0
    BaseScope = 0
    Exploitability = 0
    Impact = 0
    BaseScore = 0
    ISS = 0

    def __init__(self):
        super().__init__(root)

        self.title("Расчет базовых показателей")
        self.geometry("660x600")
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        self.base_score_entry = tk.Entry(self)
        self.base_score_entry.place(x=480, y=10, width=150)

        # Кнопки вектора атаки
        attack_vector_label = tk.Label(self, text="Вектор атаки")
        attack_vector_label.place(x=30, y=20, width=150)
        self.avn_btn = tk.Button(self, text="Сеть", bg="white", command=lambda: self.attack_vector(1))
        self.avn_btn.place(x=30, y=45, width=150)
        self.avn_btn_ttp = CreateToolTip(self.avn_btn, "Уязвимый компонент привязан к сетевому стеку\n и набор "
                                                       "возможных атак выходят за рамки других параметров,\n "
                                                       "вплоть до всего интернета.")

        self.ava_btn = tk.Button(self, text="Смежная сеть", bg="white", command=lambda: self.attack_vector(2))
        self.ava_btn.place(x=180, y=45, width=150)
        self.ava_btn_ttp = CreateToolTip(self.ava_btn, "Уязвимый компонент привязан к сетевому стеку,\n"
                                                       "но атака ограничивается на протокольном уровне\n "
                                                       "логически соседней сети.")

        self.avl_btn = tk.Button(self, text="Локальная сеть", bg="white", command=lambda: self.attack_vector(3))
        self.avl_btn.place(x=330, y=45, width=150)
        self.avl_btn_ttp = CreateToolTip(self.avl_btn, "Уязвимый компонент не привязан к сетевому стеку,\n "
                                                       "а путь злоумышленника проходит через возможности\n "
                                                       "чтения/записи/исполнения.")

        self.avp_btn = tk.Button(self, text="Физический доступ", bg="white", command=lambda: self.attack_vector(4))
        self.avp_btn.place(x=480, y=45, width=150)
        self.avp_btn_ttp = CreateToolTip(self.avp_btn, "Атака требует от злоумышленника физического взаимодействия\n "
                                                       "или манипулирования уязвимым компонентом.")

        # Кнопки Сложности эксплуатации уязвимости
        attack_complexity_label = tk.Label(self, text="Сложность эксплуатации")
        attack_complexity_label.place(x=30, y=85, width=150)
        self.acl_btn = tk.Button(self, text="Низкая", bg="white", command=lambda: self.attack_complexity(1))
        self.acl_btn.place(x=30, y=110, width=300)
        self.acl_btn_ttp = CreateToolTip(self.acl_btn, "Специализированных условий доступа или смягчающих\n"
                                                      "обстоятельств не существует. Атакующий может\n "
                                                      "рассчитывать на повторный успех при атаке на уязвимый\n"
                                                      " компонент.")

        self.ach_btn = tk.Button(self, text="Высокая", bg="white", command=lambda: self.attack_complexity(2))
        self.ach_btn.place(x=330, y=110, width=300)
        self.ach_btn_ttp = CreateToolTip(self.ach_btn, "Успешная атака требует от злоумышленника вложить\n"
                                                 "определенный измеримый объем усилий в подготовку\n "
                                                 "или выполнение атаки на уязвимый компонент, прежде\n "
                                                 "чем можно будет ожидать успешной атаки.")

        # Кнопки влияния на другие компоненты системы ( Scope )
        scope_label = tk.Label(self, text="Влияние")
        scope_label.place(x=30, y=150, width=150)
        self.sc_btn = tk.Button(self, text="Оказывает", bg="white", command=lambda: self.scope(1))
        self.sc_btn.place(x=30, y=175, width=300)
        self.sc_btn_ttp = CreateToolTip(self.sc_btn, "Эксплуатируемая уязвимость влияет на другие компоненты\n системы")

        self.su_btn = tk.Button(self, text="Не оказывает", bg="white", command=lambda: self.scope(2))
        self.su_btn.place(x=330, y=175, width=300)
        self.su_btn_ttp = CreateToolTip(self.su_btn, "Эксплуатируемая уязвимость не влияет на другие компоненты\n"
                                                     " системы.")

        # Кнопки требуемого уровня привилегий
        privileges_required_label = tk.Label(self, text="Уровень привилегий")
        privileges_required_label.place(x=30, y=215, width=150)
        self.prn_btn = tk.Button(self, text="Не требуется", bg="white", command=lambda: self.privileges_required(1))
        self.prn_btn.place(x=30, y=240, width=200)
        self.prn_btn_ttp = CreateToolTip(self.prn_btn,"Атакующий является неавторизованным до начала атаки и\n"
                                                      " поэтому не требует никакого доступа к настройкам или файлам\n"
                                                      " уязвимой системы для проведения атаки")

        self.prl_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.privileges_required(2))
        self.prl_btn.place(x=230, y=240, width=200)
        self.prl_btn_ttp = CreateToolTip(self.prl_btn, "Атакующий требует привилегий, которые обеспечивают\n"
                                                     " основные пользовательские возможности, которые обычно\n"
                                                     " могут влиять только на настройки и файлы, принадлежащие\n"
                                                     " пользователю.")

        self.prh_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.privileges_required(3))
        self.prh_btn.place(x=430, y=240, width=200)
        self.prh_btn_ttp = CreateToolTip(self.prh_btn, "Атакующему требуются привилегии, обеспечивающие\n"
                                                       " значительный (например, административный) контроль над\n"
                                                       " уязвимым компонентом, позволяющий получить доступ к\n"
                                                       " настройкам и файлам всего компонента).")

        # Взаимодействие с пользователем ( User Interaction )
        user_interaction_label = tk.Label(self, text="Взаимодействие с пользователем")
        user_interaction_label.place(x=30, y=280, width=200)
        self.uin_btn = tk.Button(self, text="Не требуется", bg="white", command=lambda: self.user_interaction(1))
        self.uin_btn.place(x=30, y=305, width=300)
        self.uin_btn_ttp = CreateToolTip(self.uin_btn, "Уязвимая система может быть использована без\n"
                                                       " взаимодействия с пользователем.")

        self.uir_btn = tk.Button(self, text="Требуется", bg="white", command=lambda: self.user_interaction(2))
        self.uir_btn.place(x=330, y=305, width=300)
        self.uir_btn_ttp = CreateToolTip(self.uir_btn, "Успешная эксплуатация этой уязвимости требует от\n"
                                                       " пользователя принятия определенных мер, прежде чем эта\n"
                                                       " уязвимость может быть использована.")

        # Влияние на конфиденциальность ( confidentiality )
        confidentiality_label = tk.Label(self, text="Влияние на конфиденциальность")
        confidentiality_label.place(x=30, y=345, width=200)
        self.cn_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.confidentiality(1))
        self.cn_btn.place(x=30, y=370, width=200)
        self.cn_btn_ttp = CreateToolTip(self.cn_btn, "Внутри затронутого компонента не происходит\n"
                                                     "потери конфиденциальности.")

        self.cl_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.confidentiality(2))
        self.cl_btn.place(x=230, y=370, width=200)
        self.cl_btn_ttp = CreateToolTip(self.cl_btn, "Существует небольшая потеря конфиденциальности. Доступ к\n"
                                                     " некоторой ограниченной информации получен, но\n"
                                                     " злоумышленник не имеет контроля над тем, какая информация\n"
                                                     " получена, или размер или вид потерь ограничен.")

        self.ch_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.confidentiality(3))
        self.ch_btn.place(x=430, y=370, width=200)
        self.ch_btn_ttp = CreateToolTip(self.ch_btn, "Происходит полная потеря конфиденциальности, в результате\n"
                                                     " чего все ресурсы затронутого компонента разглашаются\n"
                                                     " злоумышленнику. В качестве альтернативы предоставляется\n"
                                                     " доступ только к некоторой ограниченной информации, однако\n"
                                                     " раскрытая информация оказывает прямое, серьезное воздействие. ")

        # Влияние на целостность ( Integrity )
        integrity_label = tk.Label(self, text="Влияние на целостность")
        integrity_label.place(x=30, y=410, width=200)
        self.in_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.integrity(1))
        self.in_btn.place(x=30, y=435, width=200)
        self.in_btn_ttp = CreateToolTip(self.in_btn, "Внутри компонента, подвергшегося воздействию, не происходит\n"
                                                     "потери целостности.")

        self.il_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.integrity(2))
        self.il_btn.place(x=230, y=435, width=200)
        self.il_btn_ttp = CreateToolTip(self.il_btn, "Модификация данных возможна, но злоумышленник не имеет\n"
                                                     " контроля над последствиями модификации, либо количество\n"
                                                     " модификаций ограничено. Изменение данных не оказывает\n"
                                                     " прямого, серьезного влияния на пораженный компонент.")

        self.ih_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.integrity(3))
        self.ih_btn.place(x=430, y=435, width=200)
        self.ih_btn_ttp = CreateToolTip(self.ih_btn, "Происходит полная потеря целостности или полная утрата защиты.")

        # Влияние на доступность ( Availability )
        availability_label = tk.Label(self, text="Влияние на доступность")
        availability_label.place(x=30, y=475, width=200)
        self.an_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.availability(1))
        self.an_btn.place(x=30, y=500, width=200)
        self.an_btn_ttp = CreateToolTip(self.an_btn, "Наличие компонента, подвергшегося воздействию,\n"
                                                     "не влияет на его доступность.")

        self.al_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.availability(2))
        self.al_btn.place(x=230, y=500, width=200)
        self.al_btn_ttp = CreateToolTip(self.al_btn, "Производительность снижается или возникают перебои с\n"
                                                     " доступностью ресурсов. Даже при возможности повторного\n"
                                                     " использования уязвимости злоумышленник не имеет возможности\n"
                                                     " полностью отказать в обслуживании законным пользователям. ")

        self.ah_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.availability(3))
        self.ah_btn.place(x=430, y=500, width=200)
        self.ah_btn_ttp = CreateToolTip(self.ah_btn, "Существует полная потеря доступности, в результате чего\n"
                                                     " злоумышленник может полностью отказывать в доступе к\n"
                                                     " ресурсам по затрагиваемому компоненту; эта потеря также\n"
                                                     " составляет устойчивым (пока нападающий продолжает совершать\n"
                                                     " нападение) или настойчивым (нападающий состояние сохраняется\n"
                                                     " даже после завершения атаки).")

        # Кнопка вывод
        base_score_btn = tk.Button(self, text="Считать", bg="white", command=self.base_score)
        base_score_btn.place(x=480, y=550, width=150)

        # Кнопка ок ( закрывает дочернее откно и передает в основное байзскор)
    # Методы кнопок вектора атаки
    def attack_vector(self, a):
        if a == 1:
            Base.AttackVector = 0.85
            self.avn_btn.configure(bg="#A9D0F5")
            self.ava_btn.configure(bg="white")
            self.avl_btn.configure(bg="white")
            self.avp_btn.configure(bg="white")
        elif a == 2:
            Base.AttackVector = 0.62
            self.avn_btn.configure(bg="white")
            self.ava_btn.configure(bg="#A9D0F5")
            self.avl_btn.configure(bg="white")
            self.avp_btn.configure(bg="white")
        elif a == 3:
            Base.AttackVector = 0.55
            self.avn_btn.configure(bg="white")
            self.ava_btn.configure(bg="white")
            self.avl_btn.configure(bg="#A9D0F5")
            self.avp_btn.configure(bg="white")
        elif a == 4:
            Base.AttackVector = 0.2
            self.avn_btn.configure(bg="white")
            self.ava_btn.configure(bg="white")
            self.avl_btn.configure(bg="white")
            self.avp_btn.configure(bg="#A9D0F5")

    # Параметры Сложности эксплуатации уязвимости
    def attack_complexity(self, a):
        if a == 1:
            Base.AttackComplexity = 0.77
            self.acl_btn.configure(bg="#A9D0F5")
            self.ach_btn.configure(bg="white")
        if a == 2:
            Base.AttackComplexity = 0.44
            self.acl_btn.configure(bg="white")
            self.ach_btn.configure(bg="#A9D0F5")

    # Влияние на другие компоненты системы ( Scope )
    def scope(self, a):
        if a == 1:
            Base.Scope = 1
            self.sc_btn.configure(bg="#A9D0F5")
            self.su_btn.configure(bg="white")
        if a == 2:
            Base.Scope = 0
            self.sc_btn.configure(bg="white")
            self.su_btn.configure(bg="#A9D0F5")

    # Требуемый уровень привилегий
    def privileges_required(self, a):
        if Base.Scope == 0:
            if a == 1:
                Base.PrivilegesRequired = 0.85
                self.prn_btn.configure(bg="#A9D0F5")
                self.prl_btn.configure(bg="white")
                self.prh_btn.configure(bg="white")
            elif a == 2:
                Base.PrivilegesRequired = 0.62
                self.prn_btn.configure(bg="white")
                self.prl_btn.configure(bg="#A9D0F5")
                self.prh_btn.configure(bg="white")
            elif a == 3:
                Base.PrivilegesRequired = 0.27
                self.prn_btn.configure(bg="white")
                self.prl_btn.configure(bg="white")
                self.prh_btn.configure(bg="#A9D0F5")

        elif Base.Scope == 1:
            if a == 1:
                Base.PrivilegesRequired = 0.85
                self.prn_btn.configure(bg="#A9D0F5")
                self.prl_btn.configure(bg="white")
                self.prh_btn.configure(bg="white")
            elif a == 2:
                Base.PrivilegesRequired = 0.68
                self.prn_btn.configure(bg="white")
                self.prl_btn.configure(bg="#A9D0F5")
                self.prh_btn.configure(bg="white")
            elif a == 3:
                Base.PrivilegesRequired = 0.5
                self.prn_btn.configure(bg="white")
                self.prl_btn.configure(bg="white")
                self.prh_btn.configure(bg="#A9D0F5")

    # Взаимодействие с пользователем ( User Interaction )
    def user_interaction(self, a):
        if a == 1:
            Base.UserInteraction = 0.85
            self.uin_btn.configure(bg="#A9D0F5")
            self.uir_btn.configure(bg="white")
        if a == 2:
            Base.UserInteraction = 0.62
            self.uin_btn.configure(bg="white")
            self.uir_btn.configure(bg="#A9D0F5")

    # Влияние на конфиденциальность ( Confidentiality )
    def confidentiality(self, a):
        if a == 1:
            Base.Confidentiality = 0
            self.cn_btn.configure(bg="#A9D0F5")
            self.cl_btn.configure(bg="white")
            self.ch_btn.configure(bg="white")
        elif a == 2:
            Base.Confidentiality = 0.22
            self.cn_btn.configure(bg="white")
            self.cl_btn.configure(bg="#A9D0F5")
            self.ch_btn.configure(bg="white")
        elif a == 3:
            Base.Confidentiality = 0.56
            self.cn_btn.configure(bg="white")
            self.cl_btn.configure(bg="white")
            self.ch_btn.configure(bg="#A9D0F5")

    # Влияние на целостность ( Integrity )
    def integrity(self, a):
        if a == 1:
            Base.Integrity = 0
            self.in_btn.configure(bg="#A9D0F5")
            self.il_btn.configure(bg="white")
            self.ih_btn.configure(bg="white")
        elif a == 2:
            Base.Integrity = 0.22
            self.in_btn.configure(bg="white")
            self.il_btn.configure(bg="#A9D0F5")
            self.ih_btn.configure(bg="white")
        elif a == 3:
            Base.Integrity = 0.56
            self.in_btn.configure(bg="white")
            self.il_btn.configure(bg="white")
            self.ih_btn.configure(bg="#A9D0F5")

    # Влияние на доступность ( Availability )
    def availability(self, a):
        if a == 1:
            Base.Availability = 0
            self.an_btn.configure(bg="#A9D0F5")
            self.al_btn.configure(bg="white")
            self.ah_btn.configure(bg="white")
        elif a == 2:
            Base.Availability = 0.22
            self.an_btn.configure(bg="white")
            self.al_btn.configure(bg="#A9D0F5")
            self.ah_btn.configure(bg="white")
        elif a == 3:
            Base.Availability = 0.56
            self.an_btn.configure(bg="white")
            self.al_btn.configure(bg="white")
            self.ah_btn.configure(bg="#A9D0F5")

    def base_score(self):

        # Расчет базовых показателей

        Base.Exploitability = 8.22 * Base.AttackVector * Base.AttackComplexity * Base.PrivilegesRequired\
                             * Base.UserInteraction
        print('Exploitability : ', Base.Exploitability)
        Base.Impact = 1 - ((1 - Base.Confidentiality) * (1 - Base.Integrity) * (1 - Base.Availability))
        print('Impact :', Base.Impact)
        # changed
        if Base.Scope == 1:
            Base.ISS = 7.52 * (Base.Impact - 0.029) - 3.25 * ((Base.Impact - 0.02) ** 15)
        # unchanged
        else:
            Base.ISS = 6.42 * Base.Impact

        # Impact sub score
        if Base.ISS <= 0:
            Base.BaseScore = 0

        # changed , потом не забыть округлить и выбрать меньшее
        elif Base.Scope == 1:
            Base.BaseScore = 1.08 * (Base.ISS + Base.Exploitability)
            Base.BaseScore = round(Base.BaseScore, 2)
            print('bs', Base.BaseScore)
        # unchanged , потом не забыть округлить и выбрать меньшее
        elif Base.Scope == 0:
            Base.BaseScore = Base.ISS + Base.Exploitability
            Base.BaseScore = round(Base.BaseScore, 2)
            print('bs', Base.BaseScore)

        self.base_score_entry.delete(0, tk.END)
        self.base_score_entry.insert(0, Base.BaseScore)

        # Вывод результата в мейн окно
        Main.baseentry.delete(0, tk.END)
        Main.baseentry.insert(0, Base.BaseScore)


class Temporal(tk.Toplevel):
    ExploitCodeMaturity = 1
    RemendiationLevel = 1
    ReportConfidence = 1
    TemporalScore = 0

    def __init__(self):
        super().__init__(root)
        self.title("Расчет временных показателей")
        self.geometry("810x600")
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        self.base_temporal_score_entry = tk.Entry(self)
        self.base_temporal_score_entry.place(x=630, y=10, width=150)

        # Кнопки зрелости срендств эксплуатации
        ExploitCodeMaturity_label = tk.Label(self, text="Зрелость средств эксплуатации")
        ExploitCodeMaturity_label.place(x=30, y=20, width=200)
        self.ExploitCodeMaturity_notdefined_btn = tk.Button(self, text="Не определено", bg="white",
                                                            command=lambda: self.exm(1))
        self.ExploitCodeMaturity_notdefined_btn.place(x=30, y=45, width=150)
        self.ExploitCodeMaturity_notdefined_btn_ttp = CreateToolTip(self.ExploitCodeMaturity_notdefined_btn,
                                                    "Присвоение этого значения указывает на недостаточность\n"
                                                    "информации для выбора одного из других значений и не влияет на\n"
                                                    "общую временную оценку, т.е. оказывает такое же влияние на\n"
                                                    "баллы, как и присвоение значения - высокая")

        self.ExploitCodeMaturity_high_btn = tk.Button(self, text="Высокая", bg="white",
                                                      command=lambda: self.exm(2))
        self.ExploitCodeMaturity_high_btn.place(x=180, y=45, width=150)
        self.ExploitCodeMaturity_high_btn_ttp = CreateToolTip(self.ExploitCodeMaturity_high_btn,
                                                    "Функциональный автономный код существует, или не требуется\n"
                                                    "никакого использования (ручной триггер), и детали широко\n"
                                                    "доступны. Код эксплойта работает в любой ситуации или активно\n"
                                                    "передается через автономного агента (например, червя или вируса).")

        self.ExploitCodeMaturity_functional_btn = tk.Button(self, text="Есть сценарий", bg="white",
                                                            command=lambda: self.exm(3))
        self.ExploitCodeMaturity_functional_btn.place(x=330, y=45, width=150)
        self.ExploitCodeMaturity_functional_btn_ttp = CreateToolTip(self.ExploitCodeMaturity_functional_btn,
                                                    "Функциональный код эксплойта доступен. Код работает в\n"
                                                    "большинстве ситуаций, когда уязвимость существует.")

        self.ExploitCodeMaturity_poc_btn = tk.Button(self, text="Есть PoC-код", bg="white",
                                                     command=lambda: self.exm(4))
        self.ExploitCodeMaturity_poc_btn.place(x=480, y=45, width=150)
        self.ExploitCodeMaturity_poc_btn_ttp = CreateToolTip(self.ExploitCodeMaturity_poc_btn,
                                                    "Доказательство концепции кода эксплойта доступно, или\n"
                                                    "демонстрация атаки не практична для большинства систем. Код\n"
                                                    "или техника не функционирует во всех ситуациях и может\n"
                                                    "потребовать существенной модификации со стороны\n"
                                                    "квалифицированного атакующего.")

        self.ExploitCodeMaturity_unproven_btn = tk.Button(self, text="Есть Теория", bg="white",
                                                          command=lambda: self.exm(5))
        self.ExploitCodeMaturity_unproven_btn.place(x=630, y=45, width=150)
        self.ExploitCodeMaturity_unproven_btn_ttp = CreateToolTip(self.ExploitCodeMaturity_unproven_btn,
                                                    "Код эксплойта недоступен, либо эксплойт является теоретическим.")

        # Кнопки доступность средств устранения
        RemendiationLevel_label = tk.Label(self, text="Доступность средств устранения уязвимости")
        RemendiationLevel_label.place(x=30, y=85, width=300)
        self.RemendiationLevel_notdefined_btn = tk.Button(self, text="Не определено", bg="white",
                                                          command=lambda: self.remendiation_level(1))
        self.RemendiationLevel_notdefined_btn.place(x=30, y=110, width=150)
        self.RemendiationLevel_notdefined_btn_ttp = CreateToolTip(self.RemendiationLevel_notdefined_btn,
                                                    "Присвоение этого значения указывает на недостаточность\n"
                                                    "информации для выбора одного из других значений и не влияет на\n"
                                                    " общую временную оценку, т.е. влияет на оценку так же, как и\n"
                                                    " присвоение не доступно.")

        self.RemendiationLevel_unavailable_btn = tk.Button(self, text="Не доступно", bg="white",
                                                           command=lambda: self.remendiation_level(2))
        self.RemendiationLevel_unavailable_btn.place(x=180, y=110, width=150)
        self.RemendiationLevel_unavailable_btn_ttp = CreateToolTip(self.RemendiationLevel_unavailable_btn,
                                                    "Решения либо нет, либо его невозможно применить.")

        self.RemendiationLevel_workaround_btn = tk.Button(self, text="Есть рекомендации", bg="white",
                                                          command=lambda: self.remendiation_level(3))
        self.RemendiationLevel_workaround_btn.place(x=330, y=110, width=150)
        self.RemendiationLevel_workaround_btn_ttp = CreateToolTip(self.RemendiationLevel_workaround_btn,
                                                    "Существует неофициальное, не связанное с поставщиками\n"
                                                    "решение. В некоторых случаях пользователи затронутой\n"
                                                    "технологии создадут свой собственный патч или обеспечат шаги\n"
                                                    "для работы или иного уменьшения уязвимости.")

        self.RemendiationLevel_temporaryfix_btn = tk.Button(self, text="Временное", bg="white",
                                                            command=lambda: self.remendiation_level(4))
        self.RemendiationLevel_temporaryfix_btn.place(x=480, y=110, width=150)
        self.RemendiationLevel_temporaryfix_btn_ttp = CreateToolTip(self.RemendiationLevel_temporaryfix_btn,
                                                    "Есть официальное, но временное решение. Сюда относятся случаи,\n"
                                                    " когда поставщик выпускает временное исправление, инструмент\n"
                                                    "или обходной путь.")

        self.RemendiationLevel_officialfix_btn = tk.Button(self, text="Официальное", bg="white",
                                                           command=lambda: self.remendiation_level(5))
        self.RemendiationLevel_officialfix_btn.place(x=630, y=110, width=150)
        self.RemendiationLevel_officialfix_btn_ttp = CreateToolTip(self.RemendiationLevel_officialfix_btn,
                                                    "Доступно полное решение от производителя. Продавец либо\n"
                                                    "выпустил официальный патч, либо доступно обновление.")

        # Кнопки степени доверия к информации
        ReportConfidence_label = tk.Label(self, text="Степень доверия к информации об уязвимости")
        ReportConfidence_label.place(x=30, y=150, width=300)
        self.ReportConfidence_notdefined_btn = tk.Button(self, text="Не определена", bg="white",
                                                         command=lambda: self.report_confidence(1))
        self.ReportConfidence_notdefined_btn.place(x=30, y=175, width=150)
        self.ReportConfidence_notdefined_btn_ttp = CreateToolTip(self.ReportConfidence_notdefined_btn,
                                                        "Присвоение этого значения указывает на недостаточность\n"
                                                        "информации для выбора одного из других значений и не влияет\n"
                                                        "на общий Временную оценку, т.е. влияет на балл так же, как и\n"
                                                        "присвоение параметра «Подтверждено».")

        self.ReportConfidence_unknown_btn = tk.Button(self, text="Не известно", bg="white",
                                                      command=lambda: self.report_confidence(2))
        self.ReportConfidence_unknown_btn.place(x=180, y=175, width=150)
        self.ReportConfidence_unknown_btn_ttp = CreateToolTip(self.ReportConfidence_unknown_btn,
                                                        "Имеются сообщения о воздействии, указывающие на наличие\n"
                                                        " уязвимости, отчеты указывают на то, что причина уязвимости\n"
                                                        " неизвестна, или отчеты могут различаться в зависимости от\n"
                                                        " причины или последствий уязвимости")

        self.ReportConfidence_reasonable_btn = tk.Button(self, text="Есть дост. отчеты", bg="white",
                                                         command=lambda: self.report_confidence(3))
        self.ReportConfidence_reasonable_btn.place(x=330, y=175, width=150)
        self.ReportConfidence_reasonable_btn_ttp = CreateToolTip(self.ReportConfidence_reasonable_btn,
                                                        "Публикуются существенные детали, но исследователи либо не\n"
                                                        " имеют оснований в них верить или не имеют доступ к исходному\n"
                                                        " коду, чтобы полностью подтвердить все взаимодействия,\n"
                                                        " которые могут привести к результату.  ")

        self.ReportConfidence_confirmed_btn = tk.Button(self, text="Подтверждена", bg="white",
                                                        command=lambda: self.report_confidence(4))
        self.ReportConfidence_confirmed_btn.place(x=480, y=175, width=150)
        self.ReportConfidence_confirmed_btn_ttp = CreateToolTip(self.ReportConfidence_confirmed_btn,
                                                        "Существуют подробные отчеты, или возможно функциональное\n"
                                                        " воспроизведение (это может быть обеспечено функциональными\n"
                                                        " эксплойтами). Исходный код доступен для независимой проверки\n"
                                                        " утверждений исследования, либо автор или продавец\n"
                                                        " пораженного кода подтвердил наличие уязвимости.")

        # Кнопка подсчета временных показателей
        base_temporal_score_btn = tk.Button(self, text="Считать",  bg="white", command=self.base_temporal_score)
        base_temporal_score_btn.place(x=630, y=225, width=150)

        # Кнопки зрелости срендств эксплуатации
    def exm(self, a):
        if a == 1:
            Temporal.ExploitCodeMaturity = 1
            self.ExploitCodeMaturity_notdefined_btn.configure(bg="#A9D0F5")
            self.ExploitCodeMaturity_high_btn.configure(bg="white")
            self.ExploitCodeMaturity_functional_btn.configure(bg="white")
            self.ExploitCodeMaturity_poc_btn.configure(bg="white")
            self.ExploitCodeMaturity_unproven_btn.configure(bg="white")
        if a == 2:
            Temporal.ExploitCodeMaturity = 1
            self.ExploitCodeMaturity_notdefined_btn.configure(bg="white")
            self.ExploitCodeMaturity_high_btn.configure(bg="#A9D0F5")
            self.ExploitCodeMaturity_functional_btn.configure(bg="white")
            self.ExploitCodeMaturity_poc_btn.configure(bg="white")
            self.ExploitCodeMaturity_unproven_btn.configure(bg="white")
        elif a == 3:
            Temporal.ExploitCodeMaturity = 0.97
            self.ExploitCodeMaturity_notdefined_btn.configure(bg="white")
            self.ExploitCodeMaturity_high_btn.configure(bg="white")
            self.ExploitCodeMaturity_functional_btn.configure(bg="#A9D0F5")
            self.ExploitCodeMaturity_poc_btn.configure(bg="white")
            self.ExploitCodeMaturity_unproven_btn.configure(bg="white")
        elif a == 4:
            Temporal.ExploitCodeMaturity = 0.94
            self.ExploitCodeMaturity_notdefined_btn.configure(bg="white")
            self.ExploitCodeMaturity_high_btn.configure(bg="white")
            self.ExploitCodeMaturity_functional_btn.configure(bg="white")
            self.ExploitCodeMaturity_poc_btn.configure(bg="#A9D0F5")
            self.ExploitCodeMaturity_unproven_btn.configure(bg="white")
        elif a == 5:
            Temporal.ExploitCodeMaturity = 0.91
            self.ExploitCodeMaturity_notdefined_btn.configure(bg="white")
            self.ExploitCodeMaturity_high_btn.configure(bg="white")
            self.ExploitCodeMaturity_functional_btn.configure(bg="white")
            self.ExploitCodeMaturity_poc_btn.configure(bg="white")
            self.ExploitCodeMaturity_unproven_btn.configure(bg="#A9D0F5")

    # Доступные средства устранения уязвимости (Remediation Level)
    def remendiation_level(self, a):
        if a == 1:
            Temporal.RemendiationLevel = 1
            self.RemendiationLevel_notdefined_btn.configure(bg="#A9D0F5")
            self.RemendiationLevel_unavailable_btn.configure(bg="white")
            self.RemendiationLevel_workaround_btn.configure(bg="white")
            self.RemendiationLevel_temporaryfix_btn.configure(bg="white")
            self.RemendiationLevel_officialfix_btn.configure(bg="white")
        if a == 2:
            Temporal.RemendiationLevel = 1
            self.RemendiationLevel_notdefined_btn.configure(bg="white")
            self.RemendiationLevel_unavailable_btn.configure(bg="#A9D0F5")
            self.RemendiationLevel_workaround_btn.configure(bg="white")
            self.RemendiationLevel_temporaryfix_btn.configure(bg="white")
            self.RemendiationLevel_officialfix_btn.configure(bg="white")
        elif a == 3:
            Temporal.RemendiationLevel = 0.97
            self.RemendiationLevel_notdefined_btn.configure(bg="white")
            self.RemendiationLevel_unavailable_btn.configure(bg="white")
            self.RemendiationLevel_workaround_btn.configure(bg="#A9D0F5")
            self.RemendiationLevel_temporaryfix_btn.configure(bg="white")
            self.RemendiationLevel_officialfix_btn.configure(bg="white")
        elif a == 4:
            Temporal.RemendiationLevel = 0.96
            self.RemendiationLevel_notdefined_btn.configure(bg="white")
            self.RemendiationLevel_unavailable_btn.configure(bg="white")
            self.RemendiationLevel_workaround_btn.configure(bg="white")
            self.RemendiationLevel_temporaryfix_btn.configure(bg="#A9D0F5")
            self.RemendiationLevel_officialfix_btn.configure(bg="white")
        elif a == 5:
            Temporal.RemendiationLevel = 0.95
            self.RemendiationLevel_notdefined_btn.configure(bg="white")
            self.RemendiationLevel_unavailable_btn.configure(bg="white")
            self.RemendiationLevel_workaround_btn.configure(bg="white")
            self.RemendiationLevel_temporaryfix_btn.configure(bg="white")
            self.RemendiationLevel_officialfix_btn.configure(bg="#A9D0F5")

    # Степень доверия к информации об уязвимости (Report Confidence)
    def report_confidence(self, a):
        if a == 1:
            Temporal.ReportConfidence = 1
            self.ReportConfidence_notdefined_btn.configure(bg="#A9D0F5")
            self.ReportConfidence_unknown_btn.configure(bg="white")
            self.ReportConfidence_reasonable_btn.configure(bg="white")
            self.ReportConfidence_confirmed_btn.configure(bg="white")
        if a == 2:
            Temporal.ReportConfidence = 1
            self.ReportConfidence_notdefined_btn.configure(bg="white")
            self.ReportConfidence_unknown_btn.configure(bg="#A9D0F5")
            self.ReportConfidence_reasonable_btn.configure(bg="white")
            self.ReportConfidence_confirmed_btn.configure(bg="white")
        elif a == 3:
            Temporal.ReportConfidence = 0.96
            self.ReportConfidence_notdefined_btn.configure(bg="white")
            self.ReportConfidence_unknown_btn.configure(bg="white")
            self.ReportConfidence_reasonable_btn.configure(bg="#A9D0F5")
            self.ReportConfidence_confirmed_btn.configure(bg="white")

        elif a == 4:
            Temporal.ReportConfidence = 0.92
            self.ReportConfidence_notdefined_btn.configure(bg="white")
            self.ReportConfidence_unknown_btn.configure(bg="white")
            self.ReportConfidence_reasonable_btn.configure(bg="white")
            self.ReportConfidence_confirmed_btn.configure(bg="#A9D0F5")

# вывод временной метрики
    def base_temporal_score(self):
        Temporal.TemporalScore = Base.BaseScore * Temporal.ExploitCodeMaturity * Temporal.RemendiationLevel\
                                 * Temporal.ReportConfidence
        Temporal.TemporalScore = round(Temporal.TemporalScore, 2)
        self.base_temporal_score_entry.delete(0, tk.END)
        self.base_temporal_score_entry.insert(0, Temporal.TemporalScore)
        Main.temporalentry.delete(0, tk.END)
        Main.temporalentry.insert(0, Temporal.TemporalScore)


class Environmental(tk.Toplevel):

    ModifiedScope = 0
    ModifiedAttackComplexity = 0
    ModifiedAttackVector = 0
    ModifiedUserInteraction = 0
    ModifiedPrivilegesRequired = 0
    ModifiedConfidentiality = 0
    ModifiedIntegrity = 0
    ModifiedAvailability = 0
    ConfidentialityRequirement = 0
    IntegrityRequirement = 0
    AvailabilityRequirement = 0
    ModifiedExploitability = 0
    ModifiedImpact = 0
    EnvironmentalScore = 0
    MISS = 0  # Modified Impact Sub Score

    def __init__(self):
        super().__init__(root)

        self.title("Расчет контекстных показателей")
        self.geometry("1330x600")
        self.resizable(False, False)

        self.grab_set()
        self.focus_set()

        self.env_score_entry = tk.Entry(self)
        self.env_score_entry.place(x=1150, y=10, width=150)

        # Кнопки скорректированного вектора атаки
        env_attack_vector_label = tk.Label(self, text="Скорректированный вектор атаки")
        env_attack_vector_label.place(x=30, y=20, width=350)
        self.env_avn_btn = tk.Button(self, text="Сеть", bg="white", command=lambda: self.env_attack_vector(1))
        self.env_avn_btn.place(x=30, y=45, width=150)
        self.env_ava_btn = tk.Button(self, text="Смежная сеть", bg="white", command=lambda: self.env_attack_vector(2))
        self.env_ava_btn.place(x=180, y=45, width=150)
        self.env_avl_btn = tk.Button(self, text="Локальная сеть", bg="white", command=lambda: self.env_attack_vector(3))
        self.env_avl_btn.place(x=330, y=45, width=150)
        self.env_avp_btn = tk.Button(self, text="Физический доступ", bg="white",
                                     command=lambda: self.env_attack_vector(4))
        self.env_avp_btn.place(x=480, y=45, width=150)

        # Кнопки скорректированной сложности эксплуатации уязвимости
        env_attack_complexity_label = tk.Label(self, text="Скорректированная сложность эксплуатации")
        env_attack_complexity_label.place(x=30, y=85, width=350)
        self.env_acl_btn = tk.Button(self, text="Низкая", bg="white", command=lambda: self.env_attack_complexity(1))
        self.env_acl_btn.place(x=30, y=110, width=300)
        self.env_ach_btn = tk.Button(self, text="Высокая", bg="white", command=lambda: self.env_attack_complexity(2))
        self.env_ach_btn.place(x=330, y=110, width=300)

        # Кнопки скорректированного влияния на другие компоненты системы ( Scope )
        env_scope_label = tk.Label(self, text="Скорректированное влияние на другие компоненты системы")
        env_scope_label.place(x=30, y=150, width=350)
        self.env_sc_btn = tk.Button(self, text="Оказывает", bg="white", command=lambda: self.env_scope(1))
        self.env_sc_btn.place(x=30, y=175, width=300)
        self.env_su_btn = tk.Button(self, text="Не оказывает", bg="white", command=lambda: self.env_scope(2))
        self.env_su_btn.place(x=330, y=175, width=300)

        # Кнопки скорректированного требуемого уровня привилегий
        env_privileges_required_label = tk.Label(self, text="Скорректированный уровень привилегий")
        env_privileges_required_label.place(x=30, y=215, width=350)
        self.env_prn_btn = tk.Button(self, text="Не требуется", bg="white", command=lambda: self.env_privileges(1))
        self.env_prn_btn.place(x=30, y=240, width=200)
        self.env_prl_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.env_privileges(2))
        self.env_prl_btn.place(x=230, y=240, width=200)
        self.env_prh_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.env_privileges(3))
        self.env_prh_btn.place(x=430, y=240, width=200)

        # Кнопки скорректированного взаимодействие с пользователем ( User Interaction )
        env_user_interaction_label = tk.Label(self, text="Скорректированное взаимодействие с пользователем")
        env_user_interaction_label.place(x=30, y=280, width=350)
        self.env_uin_btn = tk.Button(self, text="Не требуется", bg="white",
                                     command=lambda: self.env_user_interaction(1))
        self.env_uin_btn.place(x=30, y=305, width=300)
        self.env_uir_btn = tk.Button(self, text="Требуется", bg="white", command=lambda: self.env_user_interaction(2))
        self.env_uir_btn.place(x=330, y=305, width=300)

        # Кнопки скорректированного влияния на конфиденциальность ( confidentiality )
        env_confidentiality_label = tk.Label(self, text="Скорректированное влияние на конфиденциальность")
        env_confidentiality_label.place(x=30, y=345, width=350)
        self.env_cn_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.env_confidentiality(1))
        self.env_cn_btn.place(x=30, y=370, width=200)
        self.env_cl_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.env_confidentiality(2))
        self.env_cl_btn.place(x=230, y=370, width=200)
        self.env_ch_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.env_confidentiality(3))
        self.env_ch_btn.place(x=430, y=370, width=200)

        # Кнопки скорректированного влияния на целостность ( Integrity )
        env_integrity_label = tk.Label(self, text="Скорректированное влияние на целостность")
        env_integrity_label.place(x=30, y=410, width=350)
        self.env_in_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.env_integrity(1))
        self.env_in_btn.place(x=30, y=435, width=200)
        self.env_il_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.env_integrity(2))
        self.env_il_btn.place(x=230, y=435, width=200)
        self.env_ih_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.env_integrity(3))
        self.env_ih_btn.place(x=430, y=435, width=200)

        # Кнопки скорректированного влияния на доступность ( Availability )
        env_availability_label = tk.Label(self, text="Скорректированное влияние на доступность")
        env_availability_label.place(x=30, y=475, width=350)
        self.env_an_btn = tk.Button(self, text="Нет", bg="white", command=lambda: self.env_availability(1))
        self.env_an_btn.place(x=30, y=500, width=200)
        self.env_al_btn = tk.Button(self, text="Низкий", bg="white", command=lambda: self.env_availability(2))
        self.env_al_btn.place(x=230, y=500, width=200)
        self.env_ah_btn = tk.Button(self, text="Высокий", bg="white", command=lambda: self.env_availability(3))
        self.env_ah_btn.place(x=430, y=500, width=200)

        # Требования к безопасности(Security Requirement)
        # Кнопки требований к конфиденциальности(Confidentiality Requirement)
        cr_label = tk.Label(self, text="Требования к конфиденциальности")
        cr_label.place(x=700, y=20, width=350)
        self.cr_btn_nd = tk.Button(self, text="Не определены", bg="white", command=lambda: self.cr(1))
        self.cr_btn_nd.place(x=700, y=45, width=150)
        self.cr_btn_nd_ttp = CreateToolTip(self.cr_btn_nd, "Присвоение этого значения указывает на недостаточность\n"
                                                        "информации для выбора одного из других значений и не влияет\n"
                                                        " на общий контекстный показатель, т.е. оказывает такое же\n"
                                                        " влияние на оценку, как и присвоение «Средний».")

        self.cr_btn_low = tk.Button(self, text="Низкие", bg="white", command=lambda: self.cr(2))
        self.cr_btn_low.place(x=850, y=45, width=150)
        self.cr_btn_low_ttp = CreateToolTip(self.cr_btn_low, "Потеря конфиденциальности, скорее всего, окажет лишь\n"
                                                             " ограниченное негативное воздействие на организацию или\n"
                                                             " отдельных лиц, связанных с организацией (например,\n"
                                                             " сотрудников, клиентов).")

        self.cr_btn_mid = tk.Button(self, text="Средние", bg="white", command=lambda: self.cr(3))
        self.cr_btn_mid.place(x=1000, y=45, width=150)
        self.cr_btn_mid_ttp = CreateToolTip(self.cr_btn_mid, "Потеря конфиденциальности может оказать серьезное\n"
                                                             " негативное влияние на организацию или отдельных лиц,\n"
                                                             " связанных с организацией (например, сотрудников,\n"
                                                             " клиентов).")

        self.cr_btn_high = tk.Button(self, text="Высокие", bg="white", command=lambda: self.cr(4))
        self.cr_btn_high.place(x=1150, y=45, width=150)
        self.cr_btn_high_ttp = CreateToolTip(self.cr_btn_high, "Потеря конфиденциальности может иметь катастрофически\n"
                                                               "негативные последствия для организации или отдельных\n"
                                                               "лиц, связанных с организацией (например,\n"
                                                               "сотрудников, клиентов).")

        # Кнопки требований к целостности(Integrity Requirement)
        ir_label = tk.Label(self, text="Требования к целостности")
        ir_label.place(x=700, y=85, width=350)
        self.ir_btn_nd = tk.Button(self, text="Не определены", bg="white", command=lambda: self.ir(1))
        self.ir_btn_nd.place(x=700, y=110, width=150)
        self.ir_btn_nd_ttp = CreateToolTip(self.ir_btn_nd, "Присвоение этого значения указывает на недостаточность\n"
                                                        "информации для выбора одного из других значений и не влияет\n"
                                                        "на общий контекстный показатель, т.е. оказывает такое же\n"
                                                        "влияние на оценку, как и присвоение «Средний».")

        self.ir_btn_low = tk.Button(self, text="Низкие", bg="white", command=lambda: self.ir(2))
        self.ir_btn_low.place(x=850, y=110, width=150)
        self.ir_btn_low_ttp = CreateToolTip(self.ir_btn_low, "Потеря целостности, скорее всего, окажет лишь\n"
                                                             " ограниченное негативное воздействие на организацию или\n"
                                                             " отдельных лиц, связанных с организацией (например,\n"
                                                             " сотрудников, клиентов).")

        self.ir_btn_mid = tk.Button(self, text="Средние", bg="white", command=lambda: self.ir(3))
        self.ir_btn_mid.place(x=1000, y=110, width=150)
        self.ir_btn_mid_ttp = CreateToolTip(self.ir_btn_mid, "Потеря целостности может оказать серьезное\n"
                                                             " негативное влияние на организацию или отдельных лиц,\n"
                                                             " связанных с организацией (например, сотрудников,\n"
                                                             " клиентов).")

        self.ir_btn_high = tk.Button(self, text="Высокие", bg="white", command=lambda: self.ir(4))
        self.ir_btn_high.place(x=1150, y=110, width=150)
        self.ir_btn_high_ttp = CreateToolTip(self.ir_btn_high, "Потеря целостности может иметь катастрофически\n"
                                                               "негативные последствия для организации или отдельных\n"
                                                               "лиц, связанных с организацией (например,\n"
                                                               "сотрудников, клиентов).")

        # Кнопки требований к доступности( Availability Requirement)
        ar_label = tk.Label(self, text="Требования к доступности")
        ar_label.place(x=700, y=150, width=350)
        self.ar_btn_nd = tk.Button(self, text="Не определены", bg="white", command=lambda: self.ar(1))
        self.ar_btn_nd.place(x=700, y=175, width=150)
        self.ar_btn_nd_ttp = CreateToolTip(self.ar_btn_nd, "Присвоение этого значения указывает на недостаточность\n"
                                                        "информации для выбора одного из других значений и не влияет\n"
                                                        "на общий контекстный показатель, т.е. оказывает такое же\n"
                                                        "влияние на оценку, как и присвоение «Средний».")

        self.ar_btn_low = tk.Button(self, text="Низкие", bg="white", command=lambda: self.ar(2))
        self.ar_btn_low.place(x=850, y=175, width=150)
        self.ar_btn_low_ttp = CreateToolTip(self.ar_btn_low, "Потеря доступности, скорее всего, окажет лишь\n"
                                                             " ограниченное негативное воздействие на организацию или\n"
                                                             " отдельных лиц, связанных с организацией (например,\n"
                                                             " сотрудников, клиентов).")

        self.ar_btn_mid = tk.Button(self, text="Средние", bg="white", command=lambda: self.ar(3))
        self.ar_btn_mid.place(x=1000, y=175, width=150)
        self.ar_btn_mid_ttp = CreateToolTip(self.ar_btn_mid, "Потеря доступности может оказать серьезное\n"
                                                             " негативное влияние на организацию или отдельных лиц,\n"
                                                             " связанных с организацией (например, сотрудников,\n"
                                                             " клиентов).")

        self.ar_btn_high = tk.Button(self, text="Высокие", bg="white", command=lambda: self.ar(4))
        self.ar_btn_high.place(x=1150, y=175, width=150)
        self.ar_btn_high_ttp = CreateToolTip(self.ar_btn_high, "Потеря доступности может иметь катастрофически\n"
                                                               "негативные последствия для организации или отдельных\n"
                                                               "лиц, связанных с организацией (например,\n"
                                                               "сотрудников, клиентов).")

        # Кнопка вывод
        env_score_btn = tk.Button(self, text="Считать", bg="white", command=self.enviromental_score)
        env_score_btn.place(x=1150, y=550, width=150)

        # Кнопка ок ( закрывает дочернее откно и передает в основное байзскор)

        # Методы кнопок вектора атаки

    def env_attack_vector(self, a):
        if a == 1:
            Environmental.ModifiedAttackVector = 0.85
            self.env_avn_btn.configure(bg="#A9D0F5")
            self.env_ava_btn.configure(bg="white")
            self.env_avl_btn.configure(bg="white")
            self.env_avp_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedAttackVector = 0.62
            self.env_avn_btn.configure(bg="white")
            self.env_ava_btn.configure(bg="#A9D0F5")
            self.env_avl_btn.configure(bg="white")
            self.env_avp_btn.configure(bg="white")
        elif a == 3:
            Environmental.ModifiedAttackVector = 0.55
            self.env_avn_btn.configure(bg="white")
            self.env_ava_btn.configure(bg="white")
            self.env_avl_btn.configure(bg="#A9D0F5")
            self.env_avp_btn.configure(bg="white")
        elif a == 4:
            Environmental.ModifiedAttackVector = 0.2
            self.env_avn_btn.configure(bg="white")
            self.env_ava_btn.configure(bg="white")
            self.env_avl_btn.configure(bg="white")
            self.env_avp_btn.configure(bg="#A9D0F5")

        # Параметры Сложности эксплуатации уязвимости
    def env_attack_complexity(self, a):
        if a == 1:
            Environmental.ModifiedAttackComplexity = 0.77
            self.env_acl_btn.configure(bg="#A9D0F5")
            self.env_ach_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedAttackComplexity = 0.44
            self.env_acl_btn.configure(bg="white")
            self.env_ach_btn.configure(bg="#A9D0F5")

        # Влияние на другие компоненты системы ( Scope )
    def env_scope(self, a):
        if a == 1:
            Environmental.ModifiedScope = 1
            self.env_sc_btn.configure(bg="#A9D0F5")
            self.env_su_btn.configure(bg="white")

        if a == 2:
            Environmental.ModifiedScope = 0
            self.env_sc_btn.configure(bg="white")
            self.env_su_btn.configure(bg="#A9D0F5")

    # Требуемый уровень привилегий
    def env_privileges(self, a):
        if Environmental.ModifiedScope == 0:
            if a == 1:
                Environmental.ModifiedPrivilegesRequired = 0.85
                self.env_prn_btn.configure(bg="#A9D0F5")
                self.env_prl_btn.configure(bg="white")
                self.env_prh_btn.configure(bg="white")
            elif a == 2:
                Environmental.ModifiedPrivilegesRequired = 0.62
                self.env_prn_btn.configure(bg="white")
                self.env_prl_btn.configure(bg="#A9D0F5")
                self.env_prh_btn.configure(bg="white")
            elif a == 3:
                Environmental.ModifiedPrivilegesRequired = 0.27
                self.env_prn_btn.configure(bg="white")
                self.env_prl_btn.configure(bg="white")
                self.env_prh_btn.configure(bg="#A9D0F5")
        elif Environmental.ModifiedScope == 1:
            if a == 1:
                Environmental.ModifiedPrivilegesRequired = 0.85
                self.env_prn_btn.configure(bg="#A9D0F5")
                self.env_prl_btn.configure(bg="white")
                self.env_prh_btn.configure(bg="white")
            elif a == 2:
                Environmental.ModifiedPrivilegesRequired = 0.68
                self.env_prn_btn.configure(bg="white")
                self.env_prl_btn.configure(bg="#A9D0F5")
                self.env_prh_btn.configure(bg="white")
            elif a == 3:
                Environmental.ModifiedPrivilegesRequired = 0.5
                self.env_prn_btn.configure(bg="white")
                self.env_prl_btn.configure(bg="white")
                self.env_prh_btn.configure(bg="#A9D0F5")

        # Взаимодействие с пользователем ( User Interaction )
    def env_user_interaction(self, a):
        if a == 1:
            Environmental.ModifiedUserInteraction = 0.85
            self.env_uin_btn.configure(bg="#A9D0F5")
            self.env_uir_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedUserInteraction = 0.62
            self.env_uin_btn.configure(bg="white")
            self.env_uir_btn.configure(bg="#A9D0F5")

        # Влияние на конфиденциальность ( Confidentiality )
    def env_confidentiality(self, a):
        if a == 1:
            Environmental.ModifiedConfidentiality = 0
            self.env_cn_btn.configure(bg="#A9D0F5")
            self.env_cl_btn.configure(bg="white")
            self.env_ch_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedConfidentiality = 0.22
            self.env_cn_btn.configure(bg="white")
            self.env_cl_btn.configure(bg="#A9D0F5")
            self.env_ch_btn.configure(bg="white")
        elif a == 3:
            Environmental.ModifiedConfidentiality = 0.56
            self.env_cn_btn.configure(bg="white")
            self.env_cl_btn.configure(bg="white")
            self.env_ch_btn.configure(bg="#A9D0F5")

        # Влияние на целостность ( Integrity )
    def env_integrity(self, a):
        if a == 1:
            Environmental.ModifiedIntegrity = 0
            self.env_in_btn.configure(bg="#A9D0F5")
            self.env_il_btn.configure(bg="white")
            self.env_ih_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedIntegrity = 0.22
            self.env_in_btn.configure(bg="white")
            self.env_il_btn.configure(bg="#A9D0F5")
            self.env_ih_btn.configure(bg="white")
        elif a == 3:
            Environmental.ModifiedIntegrity = 0.56
            self.env_in_btn.configure(bg="white")
            self.env_il_btn.configure(bg="white")
            self.env_ih_btn.configure(bg="#A9D0F5")

        # Влияние на доступность ( Availability )
    def env_availability(self, a):
        if a == 1:
            Environmental.ModifiedAvailability = 0
            self.env_an_btn.configure(bg="#A9D0F5")
            self.env_al_btn.configure(bg="white")
            self.env_ah_btn.configure(bg="white")
        elif a == 2:
            Environmental.ModifiedAvailability = 0.22
            self.env_an_btn.configure(bg="white")
            self.env_al_btn.configure(bg="#A9D0F5")
            self.env_ah_btn.configure(bg="white")
        elif a == 3:
            Environmental.ModifiedAvailability = 0.56
            self.env_an_btn.configure(bg="white")
            self.env_al_btn.configure(bg="white")
            self.env_ah_btn.configure(bg="#A9D0F5")

    # Требования к безопасности(Security Requirement)
    # Методы требований к конфиденциальности(Confidentiality Requirement)
    def cr(self, a):
        if a == 1:
            Environmental.ConfidentialityRequirement = 1
            self.cr_btn_nd.configure(bg="#A9D0F5")
            self.cr_btn_low.configure(bg="white")
            self.cr_btn_mid.configure(bg="white")
            self.cr_btn_high.configure(bg="white")
        elif a == 2:
            Environmental.ConfidentialityRequirement = 0.5
            self.cr_btn_nd.configure(bg="white")
            self.cr_btn_low.configure(bg="#A9D0F5")
            self.cr_btn_mid.configure(bg="white")
            self.cr_btn_high.configure(bg="white")
        elif a == 3:
            Environmental.ConfidentialityRequirement = 1
            self.cr_btn_nd.configure(bg="white")
            self.cr_btn_low.configure(bg="white")
            self.cr_btn_mid.configure(bg="#A9D0F5")
            self.cr_btn_high.configure(bg="white")
        elif a == 4:
            Environmental.ConfidentialityRequirement = 1.5
            self.cr_btn_nd.configure(bg="white")
            self.cr_btn_low.configure(bg="white")
            self.cr_btn_mid.configure(bg="white")
            self.cr_btn_high.configure(bg="#A9D0F5")

    # Методы требований к целостности(Integrity Requirement)
    def ir(self, a):
        if a == 1:
            Environmental.IntegrityRequirement = 1
            self.ir_btn_nd.configure(bg="#A9D0F5")
            self.ir_btn_low.configure(bg="white")
            self.ir_btn_mid.configure(bg="white")
            self.ir_btn_high.configure(bg="white")
        elif a == 2:
            Environmental.IntegrityRequirement = 0.5
            self.ir_btn_nd.configure(bg="white")
            self.ir_btn_low.configure(bg="#A9D0F5")
            self.ir_btn_mid.configure(bg="white")
            self.ir_btn_high.configure(bg="white")
        elif a == 3:
            Environmental.IntegrityRequirement = 1
            self.ir_btn_nd.configure(bg="white")
            self.ir_btn_low.configure(bg="white")
            self.ir_btn_mid.configure(bg="#A9D0F5")
            self.ir_btn_high.configure(bg="white")
        elif a == 4:
            Environmental.IntegrityRequirement = 1.5
            self.ir_btn_nd.configure(bg="white")
            self.ir_btn_low.configure(bg="white")
            self.ir_btn_mid.configure(bg="white")
            self.ir_btn_high.configure(bg="#A9D0F5")

    # Кнопки требований к доступности(Availability Requirement)
    def ar(self, a):
        if a == 1:
            Environmental.AvailabilityRequirement = 1
            self.ar_btn_nd.configure(bg="#A9D0F5")
            self.ar_btn_low.configure(bg="white")
            self.ar_btn_mid.configure(bg="white")
            self.ar_btn_high.configure(bg="white")
        elif a == 2:
            Environmental.AvailabilityRequirement = 0.5
            self.ar_btn_nd.configure(bg="white")
            self.ar_btn_low.configure(bg="#A9D0F5")
            self.ar_btn_mid.configure(bg="white")
            self.ar_btn_high.configure(bg="white")
        elif a == 3:
            Environmental.AvailabilityRequirement = 1
            self.ar_btn_nd.configure(bg="white")
            self.ar_btn_low.configure(bg="white")
            self.ar_btn_mid.configure(bg="#A9D0F5")
            self.ar_btn_high.configure(bg="white")
        elif a == 4:
            Environmental.AvailabilityRequirement = 1.5
            self.ar_btn_nd.configure(bg="white")
            self.ar_btn_low.configure(bg="white")
            self.ar_btn_mid.configure(bg="white")
            self.ar_btn_high.configure(bg="#A9D0F5")

    def enviromental_score(self):

        Environmental.MISS = 1 - ((1 - Environmental.ConfidentialityRequirement*Environmental.ModifiedConfidentiality) *
                              (1 - Environmental.IntegrityRequirement*Environmental.ModifiedIntegrity) *
                              (1 - Environmental.AvailabilityRequirement*Environmental.ModifiedAvailability))
        if Environmental.ModifiedScope == 1:
            Environmental.ModifiedImpact = 7.52 * (Environmental.MISS - 0.029) - 3.25 * \
                                           ((Environmental.MISS * 0.9731 - 0.02) ** 13)
            print('ModifiedImpact', Environmental.ModifiedImpact)
        # unchanged
        elif Environmental.ModifiedScope == 0:
            Environmental.ModifiedImpact = 6.42 * Environmental.MISS
            print('ModifiedImpact', Environmental.ModifiedImpact)
        Environmental.ModifiedExploitability = 8.22 * Environmental.ModifiedAttackVector * \
                                               Environmental.ModifiedAttackComplexity * \
                                               Environmental.ModifiedPrivilegesRequired * \
                                               Environmental.ModifiedUserInteraction
        print('ModifiedExploitability : ', Environmental.ModifiedExploitability)
        if Environmental.ModifiedImpact <= 0:
            Environmental.EnvironmentalScore = 0
        elif Environmental.ModifiedScope == 1:
            Environmental.EnvironmentalScore = 1.08 * (Environmental.ModifiedImpact +
            Environmental.ModifiedExploitability) * Temporal.ExploitCodeMaturity * \
                                Temporal.RemendiationLevel * Temporal.ReportConfidence
            Environmental.EnvironmentalScore = round(Environmental.EnvironmentalScore, 2)
            print('EnvironmentalScore', Environmental.EnvironmentalScore)
        elif Environmental.ModifiedScope == 0:
            Environmental.EnvironmentalScore = (Environmental.ModifiedImpact + Environmental.ModifiedExploitability) \
                                * Temporal.ExploitCodeMaturity * \
                                Temporal.RemendiationLevel * Temporal.ReportConfidence
            Environmental.EnvironmentalScore = round(Environmental.EnvironmentalScore, 2)
            print('EnvironmentalScore', Environmental.EnvironmentalScore)
        self.env_score_entry.delete(0, tk.END)
        self.env_score_entry.insert(0, Environmental.EnvironmentalScore)
        Main.environmentalentry.delete(0, tk.END)
        Main.environmentalentry.insert(0, Environmental.EnvironmentalScore)


class SaveForm(tk.Toplevel):

    doc = docx.Document()


    def __init__(self):
        super().__init__(root)

        self.title("Сохранить")
        self.geometry("455x190")
        self.resizable(False, False)
        self.grab_set()
        self.focus_set()

        self.label1 = tk.Label(self, text="Введите название")
        self.label1.place(x=30, y=30, width=200, height=25)
        self.nameentry1 = tk.Entry(self)
        self.nameentry1.place(x=30, y=65, width=300, height=25)
        self.label2 = tk.Label(self, text="Введите путь к директории")
        self.label2.place(x=30, y=100, width=200, height=25)
        self.nameentry2 = tk.Entry(self)
        self.nameentry2.place(x=30, y=135, width=300, height=25)
        self.button2 = tk.Button(self, text="Сохранить", bg="white", command=self.save)
        self.button2.place(x=340, y=135, width=85, height=25)

    def save(self):
        part1 = self.doc.add_heading('Отчет об уязвимости', 0)
        part2 = self.doc.add_heading('Наименование уязвимости :', 1)
        part3 = self.doc.add_paragraph(Main.mainlistbox)
        part4 = self.doc.add_heading('Описание уязвимости :', 1)
        part5 = self.doc.add_paragraph(Main.OpisanieVul)
        part6 = self.doc.add_heading('Описание вендора :', 1)
        part7 = self.doc.add_paragraph(Main.Vendor)
        part8 = self.doc.add_heading('Название ПО :', 1)
        part9 = self.doc.add_paragraph(Main.Nazvanie)
        part10 = self.doc.add_heading('Класс уязвимости :', 1)
        part11 = self.doc.add_paragraph(Main.Klass)
        part12 = self.doc.add_heading('Уровень опасности :', 1)
        part13 = self.doc.add_paragraph(Main.UrovenOpasnosti)
        part14 = self.doc.add_heading('Меры :', 1)
        part15 = self.doc.add_paragraph(Main.Meri)

        part10 = self.doc.add_heading('Базовый балл :', 1)
        Base.BaseScore = str(Base.BaseScore)
        part11 = self.doc.add_paragraph(Base.BaseScore)
        part12 = self.doc.add_heading('Временный балл :', 1)
        Temporal.TemporalScore = str(Temporal.TemporalScore)
        part13 = self.doc.add_paragraph(Temporal.TemporalScore)
        part14 = self.doc.add_heading('Контекстный балл :', 1)
        Environmental.EnvironmentalScore = str(Environmental.EnvironmentalScore)
        part15 = self.doc.add_paragraph(Environmental.EnvironmentalScore)

        name = str(self.nameentry1.get())
        path = str(self.nameentry2.get())
        SaveForm.doc.save(path + name + '.docx')


class CreateToolTip(object):
    # Создание всплывающего окна

    def __init__(self, widget, text='widget info'):
        self.widget = widget
        self.text = text
        self.widget.bind("<Enter>", self.enter)
        self.widget.bind("<Leave>", self.close)

    def enter(self, event=None):
        x = y = 0
        x, y, cx, cy = self.widget.bbox("insert")
        x += self.widget.winfo_rootx() + 25
        y += self.widget.winfo_rooty() + 20
        # создание окна поверх того что есть
        self.tw = tk.Toplevel(self.widget)
        # Описание самого окна
        self.tw.wm_overrideredirect(True)
        self.tw.wm_geometry("+%d+%d" % (x, y))
        label = tk.Label(self.tw, text=self.text, justify='left',
                       background='#A9D0F5', relief='solid', borderwidth=1,
                       font=("times", "12", "normal"))
        label.pack(ipadx=1)

    def close(self, event=None):
        if self.tw:
            self.tw.destroy()


if __name__ == "__main__":
    root = tk.Tk()
    app = Main(root)
    app.pack()
    root.title("Cvss calculator")
    root.geometry("1300x780")
    root.resizable(False, False)

    root.mainloop()
